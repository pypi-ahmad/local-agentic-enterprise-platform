import math
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models import Document, KnowledgeChunk
from app.services.model_router import ModelRouter, Workload
from app.services.ollama_client import OllamaClient


class KnowledgeService:
    """Ingests files, creates embeddings, and performs hybrid search with citations."""

    def __init__(self, router: ModelRouter | None = None, ollama: OllamaClient | None = None) -> None:
        self.router = router or ModelRouter()
        self.ollama = ollama or OllamaClient()

    async def ingest_document(
        self,
        session: AsyncSession,
        user_id: int,
        name: str,
        file_path: Path,
        mime_type: str,
    ) -> Document:
        text = self._extract_text(file_path, mime_type)
        doc = Document(
            owner_id=user_id,
            name=name,
            mime_type=mime_type,
            storage_path=str(file_path),
            extracted_text=text,
            metadata_json={"size": file_path.stat().st_size},
        )
        session.add(doc)
        await session.commit()
        await session.refresh(doc)

        chunks = self._chunk_text(text)
        embedding_model = (await self.router.select(Workload.EMBEDDING)).selected_model
        for idx, chunk in enumerate(chunks):
            embedding = await self.ollama.embed(embedding_model, chunk)
            session.add(
                KnowledgeChunk(
                    document_id=doc.id,
                    chunk_index=idx,
                    content=chunk,
                    embedding=embedding,
                    keyword_index=chunk.lower(),
                )
            )
        await session.commit()
        return doc

    async def search(self, session: AsyncSession, query: str, top_k: int = 5) -> list[dict]:
        result = await session.execute(select(KnowledgeChunk))
        chunks = list(result.scalars().all())
        if not chunks:
            return []

        embedding_model = (await self.router.select(Workload.EMBEDDING)).selected_model
        query_vec = await self.ollama.embed(embedding_model, query)
        if not query_vec:
            return []

        scored = []
        query_lower = query.lower()
        for chunk in chunks:
            vector_score = self._cosine_similarity(query_vec, chunk.embedding)
            keyword_bonus = 0.2 if query_lower in chunk.keyword_index else 0.0
            score = vector_score + keyword_bonus
            scored.append((score, chunk))

        scored.sort(key=lambda item: item[0], reverse=True)
        top = scored[:top_k]
        return [
            {
                "chunk_id": chunk.id,
                "document_id": chunk.document_id,
                "score": float(score),
                "content": chunk.content,
                "citation": f"document:{chunk.document_id}#chunk:{chunk.chunk_index}",
            }
            for score, chunk in top
        ]

    @staticmethod
    def _extract_text(path: Path, mime_type: str) -> str:
        if mime_type == "application/pdf":
            pdf = PdfReader(str(path))
            return "\n".join([page.extract_text() or "" for page in pdf.pages])
        if mime_type in {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        }:
            doc = DocxDocument(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        return path.read_text(encoding="utf-8", errors="ignore")

    @staticmethod
    def _chunk_text(text: str, chunk_size: int = 800) -> list[str]:
        normalized = " ".join(text.split())
        if not normalized:
            return []
        return [normalized[i : i + chunk_size] for i in range(0, len(normalized), chunk_size)]

    @staticmethod
    def _cosine_similarity(v1: list[float], v2: list[float]) -> float:
        if not v1 or not v2 or len(v1) != len(v2):
            return 0.0
        dot = sum(a * b for a, b in zip(v1, v2, strict=True))
        norm1 = math.sqrt(sum(a * a for a in v1))
        norm2 = math.sqrt(sum(b * b for b in v2))
        if norm1 == 0 or norm2 == 0:
            return 0.0
        return dot / (norm1 * norm2)
